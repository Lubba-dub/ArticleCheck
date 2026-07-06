"""DOCX 自动修正器 — 直接修改 Word 文档样式

对标: Paper Format Agent
"""
from __future__ import annotations
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


class DocxAutoFixer:
    """
    Word 文档自动修正

    用法:
        fixer = DocxAutoFixer()
        fixes = fixer.apply("paper.docx", [
            {"type": "font_inconsistency", "severity": "minor", ...},
        ])
    """

    def __init__(self):
        self._modify_log: List[Dict] = []

    def apply(self, file_path: str, issues: List[Dict]) -> List[Dict]:
        """
        根据问题列表自动修正

        Args:
            file_path: .docx 路径
            issues: 格式问题列表

        Returns:
            修正日志 (modify_log)
        """
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self._modify_log = []
        doc = Document(file_path)
        backup_path = file_path + ".bak"

        # 备份
        Path(file_path).rename(backup_path)
        Path(backup_path).rename(file_path)

        for issue in issues:
            fix_type = issue.get("type", "")
            handler = self._get_handler(fix_type)
            if handler:
                handler(doc, issue, file_path)

        doc.save(file_path)
        logger.info(f"DocxAutoFixer: 修正 {len(issues)} 个问题, 保存到 {file_path}")
        return self._modify_log

    def _log(self, fix_type: str, detail: str):
        entry = {"type": fix_type, "detail": detail}
        self._modify_log.append(entry)
        logger.info(f"  修正: {detail}")

    def _get_handler(self, fix_type: str):
        handlers = {
            "font_inconsistency": self._fix_font,
            "heading_skip": self._fix_heading,
            "spacing_inconsistency": self._fix_spacing,
            "margin": self._fix_margins,
            "page_number_missing": self._fix_page_numbers,
        }
        return handlers.get(fix_type)

    def _fix_font(self, doc, issue: Dict, path: str):
        from docx.shared import Pt
        changed = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name and run.font.name.lower() not in ("times new roman", "宋体", ""):
                    old = run.font.name
                    run.font.name = "Times New Roman"
                    changed += 1
        if changed:
            self._log("font_inconsistency", f"统一字体为 Times New Roman ({changed} 处)")

    def _fix_heading(self, doc, issue: Dict, path: str):
        """修复标题层级跳跃"""
        changed = 0
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                if level.isdigit() and int(level) > 3:
                    old = para.style.name
                    para.style = doc.styles["Heading 3"]
                    changed += 1
        if changed:
            self._log("heading_skip", f"调整标题层级 ({changed} 处 → Heading 3)")

    def _fix_spacing(self, doc, issue: Dict, path: str):
        """统一段落间距"""
        changed = 0
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.space_before and pf.space_before.pt > 12:
                from docx.shared import Pt
                pf.space_before = Pt(6)
                changed += 1
        if changed:
            self._log("spacing_inconsistency", f"统一段前间距为 6pt ({changed} 处)")

    def _fix_margins(self, doc, issue: Dict, path: str):
        """修复页边距"""
        for section in doc.sections:
            from docx.shared import Cm
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
        self._log("margin", "设置页边距 2.54/2.54/3.17/3.17 cm")

    def _fix_page_numbers(self, doc, issue: Dict, path: str):
        """添加页码"""
        for section in doc.sections:
            footer = section.footer
            if not footer.paragraphs or not footer.paragraphs[0].text.strip():
                para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                from docx.oxml.ns import qn
                fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
                run._element.append(fldChar1)
                instrText = run._element.makeelement(qn('w:instrText'), {})
                instrText.text = 'PAGE'
                run._element.append(instrText)
                fldChar2 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
                run._element.append(fldChar2)
        self._log("page_number_missing", "添加页码 (页脚居中)")
