"""
统一文档层 — 读写 DOCX / LaTeX / PDF / Markdown

核心思想: 所有格式统一为 InternalDoc 中间表示，再输出到目标格式。
参考: Docling (unified DoclingDocument), all2md (bidirectional AST)

InternalDoc 结构:
  - metadata: {title, authors, date, ...}
  - sections: [{heading, level, paragraphs: [text], figures: [...], tables: [...]}}]
  - references: [{id, title, authors, year, doi, ...}]
  - equations: [{id, latex, ...}]
"""
from __future__ import annotations
import json
import logging
import re
from dataclasses import dataclass, field, asdict
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    DOCX = "docx"
    LATEX = "latex"
    PDF = "pdf"
    MARKDOWN = "markdown"
    TXT = "txt"
    UNKNOWN = "unknown"


# ─── 中间表示 ─────────────────────────────────────────

@dataclass
class TableData:
    caption: str = ""
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)


@dataclass
class FigureData:
    caption: str = ""
    path: Optional[str] = None
    alt_text: str = ""


@dataclass
class ReferenceData:
    ref_id: str = ""
    title: str = ""
    authors: List[str] = field(default_factory=list)
    year: Optional[int] = None
    doi: Optional[str] = None
    journal: Optional[str] = None
    raw_text: str = ""


@dataclass
class SectionData:
    heading: str = ""
    level: int = 1
    paragraphs: List[str] = field(default_factory=list)
    figures: List[FigureData] = field(default_factory=list)
    tables: List[TableData] = field(default_factory=list)
    equations: List[str] = field(default_factory=list)


@dataclass
class InternalDoc:
    """内部统一文档表示"""
    metadata: Dict[str, Any] = field(default_factory=lambda: {
        "title": "", "authors": "", "date": "",
    })
    sections: List[SectionData] = field(default_factory=list)
    references: List[ReferenceData] = field(default_factory=list)
    abstract: str = ""
    raw_text: str = ""

    def to_markdown(self) -> str:
        lines = []
        if self.metadata.get("title"):
            lines.append(f"# {self.metadata['title']}\n")
        if self.abstract:
            lines.append(f"> {self.abstract}\n")
        for sec in self.sections:
            prefix = "#" * min(sec.level, 6)
            lines.append(f"\n{prefix} {sec.heading}\n")
            for p in sec.paragraphs:
                if p.strip():
                    lines.append(f"{p}\n")
            for eq in sec.equations:
                lines.append(f"$$ {eq} $$\n")
            for fig in sec.figures:
                lines.append(f"![{fig.caption}]({fig.path or ''})\n")
            for tbl in sec.tables:
                if tbl.headers:
                    lines.append("| " + " | ".join(tbl.headers) + " |")
                    lines.append("|" + "|".join("---" for _ in tbl.headers) + "|")
                    for row in tbl.rows:
                        lines.append("| " + " | ".join(row) + " |")
        if self.references:
            lines.append("\n## 参考文献\n")
            for i, r in enumerate(self.references, 1):
                lines.append(f"[{i}] {r.title[:80]}... ({r.year or '?'})\n")
        return "\n".join(lines)

    def to_dict(self) -> dict:
        return {
            "metadata": self.metadata,
            "sections_count": len(self.sections),
            "paragraphs": sum(len(s.paragraphs) for s in self.sections),
            "references": len(self.references),
            "tables": sum(len(s.tables) for s in self.sections),
            "figures": sum(len(s.figures) for s in self.sections),
        }

    @classmethod
    def from_markdown(cls, text: str) -> "InternalDoc":
        doc = cls(raw_text=text)
        current_section = SectionData(heading="preamble")
        for line in text.split("\n"):
            if line.startswith("# "):
                doc.metadata["title"] = line[2:].strip()
            elif line.startswith("## "):
                if current_section.heading:
                    doc.sections.append(current_section)
                current_section = SectionData(heading=line[3:].strip(), level=2)
            elif line.strip():
                current_section.paragraphs.append(line)
        if current_section.heading:
            doc.sections.append(current_section)
        return doc


# ═══════════════════════════════════════════════════════
# 读取器
# ═══════════════════════════════════════════════════════

def _detect_format(path: Path) -> DocumentFormat:
    suffix = path.suffix.lower()
    if suffix in (".docx", ".doc"):
        return DocumentFormat.DOCX
    if suffix in (".tex", ".ltx", ".cls", ".sty"):
        return DocumentFormat.LATEX
    if suffix == ".pdf":
        return DocumentFormat.PDF
    if suffix in (".md", ".markdown"):
        return DocumentFormat.MARKDOWN
    if suffix == ".txt":
        return DocumentFormat.TXT
    return DocumentFormat.UNKNOWN


class DocReader:
    """统一文档读取器"""

    @staticmethod
    def read(path: str) -> InternalDoc:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"文件不存在: {path}")
        fmt = _detect_format(p)
        logger.info(f"读取文档: {p.name} ({fmt.value})")

        if fmt == DocumentFormat.DOCX:
            return DocReader._read_docx(p)
        elif fmt == DocumentFormat.LATEX:
            return DocReader._read_latex(p)
        elif fmt == DocumentFormat.PDF:
            return DocReader._read_pdf(p)
        elif fmt == DocumentFormat.MARKDOWN:
            return InternalDoc.from_markdown(p.read_text("utf-8", errors="replace"))
        else:
            text = p.read_text("utf-8", errors="replace")
            return InternalDoc(raw_text=text)

    @staticmethod
    def _read_docx(path: Path) -> InternalDoc:
        from docx import Document
        doc = Document(str(path))
        internal = InternalDoc()
        # Meta
        core = doc.core_properties
        internal.metadata = {
            "title": core.title or path.stem,
            "author": core.author or "",
            "created": str(core.created) if core.created else "",
        }
        # Paragraphs → sections
        current = SectionData(heading="preamble")
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower() if para.style else ""
            if "heading" in style or "title" in style:
                if current.heading:
                    internal.sections.append(current)
                level = int(re.search(r'\d+', style).group()) if re.search(r'\d+', style) else 1
                current = SectionData(heading=text, level=level)
            else:
                current.paragraphs.append(text)
        if current.heading:
            internal.sections.append(current)
        internal.raw_text = "\n".join(p.text for p in doc.paragraphs)
        # References
        from article_check.references import ReferenceEngine
        try:
            refs = ReferenceEngine().extract_from_paper(str(path))
            internal.references = [
                ReferenceData(ref_id=r.ref_id, title=r.title, authors=r.authors,
                              year=r.year, doi=r.doi, raw_text=r.raw_text)
                for r in refs
            ]
        except Exception:
            pass
        return internal

    @staticmethod
    def _read_latex(path: Path) -> InternalDoc:
        text = path.read_text("utf-8", errors="replace")
        internal = InternalDoc(raw_text=text)
        # Title
        m = re.search(r'\\title\{(.+?)\}', text)
        if m: internal.metadata["title"] = m.group(1)
        m = re.search(r'\\author\{(.+?)\}', text)
        if m: internal.metadata["author"] = m.group(1)
        # Abstract
        m = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', text, re.DOTALL)
        if m: internal.abstract = m.group(1).strip()
        # Sections
        secs = re.findall(r'\\(?:section|subsection|chapter)\*?\{(.+?)\}', text)
        current = SectionData(heading="preamble")
        for sec_title in secs:
            if current.heading:
                internal.sections.append(current)
            current = SectionData(heading=sec_title, level=1)
            internal.sections.append(current)
            current = SectionData(heading="")
        if current.heading:
            internal.sections.append(current)
        # References
        bibs = re.findall(r'\\bibitem\{.+?\}\s*(.+?)(?=\\bibitem|\Z)', text, re.DOTALL)
        for i, b in enumerate(bibs):
            yr = re.search(r'(19|20)\d{2}', b)
            internal.references.append(ReferenceData(
                ref_id=f"ref_{i+1}", raw_text=b.strip()[:200],
                year=int(yr.group()) if yr else None,
            ))
        return internal

    @staticmethod
    def _read_pdf(path: Path) -> InternalDoc:
        text = ""
        try:
            import fitz
            with fitz.open(str(path)) as doc:
                text = "\n".join(page.get_text() for page in doc)
        except ImportError:
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            except ImportError:
                logger.error("PDF 解析需要 pymupdf 或 pdfplumber")
        return InternalDoc(raw_text=text)


# ═══════════════════════════════════════════════════════
# 写入器
# ═══════════════════════════════════════════════════════

class DocWriter:
    """统一文档写入器"""

    @staticmethod
    def write(internal: InternalDoc, output_path: str, fmt: Optional[DocumentFormat] = None) -> str:
        p = Path(output_path)
        if fmt is None:
            fmt = _detect_format(p)
        if fmt == DocumentFormat.UNKNOWN:
            fmt = DocumentFormat.MARKDOWN
        logger.info(f"写入文档: {p.name} ({fmt.value})")

        if fmt == DocumentFormat.MARKDOWN:
            p.write_text(internal.to_markdown(), encoding="utf-8")
        elif fmt == DocumentFormat.DOCX:
            DocWriter._write_docx(internal, p)
        elif fmt == DocumentFormat.LATEX:
            DocWriter._write_latex(internal, p)
        elif fmt == DocumentFormat.TXT:
            p.write_text(internal.raw_text or internal.to_markdown(), encoding="utf-8")
        else:
            p.write_text(internal.to_markdown(), encoding="utf-8")
        logger.info(f"文档已保存: {p} ({p.stat().st_size} bytes)")
        return str(p)

    @staticmethod
    def _write_docx(internal: InternalDoc, path: Path):
        from docx import Document
        from docx.shared import Pt, Inches
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        # Title
        if internal.metadata.get("title"):
            doc.add_heading(internal.metadata["title"], 0)
        if internal.abstract:
            doc.add_heading("Abstract", 1)
            doc.add_paragraph(internal.abstract)
        # Sections
        for sec in internal.sections:
            if sec.heading:
                doc.add_heading(sec.heading, min(sec.level, 9))
            for p in sec.paragraphs[:50]:
                if p.strip():
                    doc.add_paragraph(p.strip())
        # References
        if internal.references:
            doc.add_heading("参考文献", 1)
            for i, r in enumerate(internal.references, 1):
                doc.add_paragraph(f"[{i}] {r.title[:100]} ({r.year or '?'})")
        doc.save(str(path))

    @staticmethod
    def _write_latex(internal: InternalDoc, path: Path):
        lines = [r"\documentclass[12pt]{article}", r"\usepackage[utf8]{inputenc}",
                 r"\usepackage{amsmath,amssymb}", r"\usepackage{mathptmx}", ""]
        if internal.metadata.get("title"):
            lines.append(f"\\title{{{internal.metadata['title']}}}")
        if internal.metadata.get("author"):
            lines.append(f"\\author{{{internal.metadata['author']}}}")
        lines.extend([r"\date{}", r"\begin{document}", r"\maketitle", ""])
        if internal.abstract:
            lines.extend([r"\begin{abstract}", internal.abstract, r"\end{abstract}", ""])
        for sec in internal.sections:
            if sec.heading:
                cmd = r"\section" if sec.level <= 1 else r"\subsection"
                lines.append(f"{cmd}{{{sec.heading}}}")
            for p in sec.paragraphs[:30]:
                if p.strip():
                    lines.append(p.strip() + "\n")
            for eq in sec.equations:
                lines.append(f"\\begin{{equation}}\n{eq}\n\\end{{equation}}")
        if internal.references:
            lines.append(r"\begin{thebibliography}{99}")
            for i, r in enumerate(internal.references, 1):
                lines.append(f"\\bibitem{{{r.ref_id or f'ref_{i}'}}} {r.raw_text[:200]}")
            lines.append(r"\end{thebibliography}")
        lines.append(r"\end{document}")
        path.write_text("\n".join(lines), encoding="utf-8")


# ═══════════════════════════════════════════════════════
# 高层统一接口
# ═══════════════════════════════════════════════════════

def read_document(path: str) -> InternalDoc:
    """读取任何支持的文档格式 → InternalDoc"""
    return DocReader.read(path)


def write_document(internal: InternalDoc, output_path: str, fmt: Optional[DocumentFormat] = None) -> str:
    """InternalDoc → 任何支持的文件格式"""
    return DocWriter.write(internal, output_path, fmt)


def convert_document(input_path: str, output_path: str) -> str:
    """文档格式转换，自动检测输入/输出格式"""
    internal = read_document(input_path)
    return write_document(internal, output_path)
